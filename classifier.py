#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
classifier.py — Lõi (engine) phân loại hợp đồng dùng chung cho cả
bản chạy theo folder (classify_contracts.py) và bản web upload (app.py).

Template là TÙY CHỌN:
  - CÓ template: 1) Giống từ 70% trở lên  2) Hợp đồng trọng yếu  3) Không theo template
  - KHÔNG template: chỉ phân  1) Hợp đồng trọng yếu  2) Hợp đồng không trọng yếu
Mọi hợp đồng đều được gắn nhãn trọng yếu/không và ghi chú quy trình review.
Việc nhận diện trọng yếu dựa trên DANH SÁCH hợp đồng trọng yếu của công ty (MATERIAL_LIST),
ưu tiên model Qwen (GreenNode MaaS), fallback dò từ khóa khi không gọi được model.
"""

import os, re, json, subprocess, tempfile, difflib

# ---------- Cấu hình ----------
THRESHOLD = 70.0   # % ngưỡng "giống template"

CAT_SIM         = "Giống từ 70% trở lên"
CAT_KEY         = "Hợp đồng trọng yếu"
CAT_NONE        = "Không theo template"
CAT_NONMATERIAL = "Hợp đồng không trọng yếu"
# Có template: phân theo template (giữ nguyên logic so sánh cũ)
CATEGORIES_WITH_TEMPLATE = [CAT_SIM, CAT_KEY, CAT_NONE]
# Không upload template: chỉ phân trọng yếu / không trọng yếu
CATEGORIES_NO_TEMPLATE   = [CAT_KEY, CAT_NONMATERIAL]
CATEGORIES = CATEGORIES_WITH_TEMPLATE   # tương thích ngược

# Ghi chú quy trình theo mức độ trọng yếu
NOTE_MATERIAL    = "Bắt buộc phải có legal review"
NOTE_NONMATERIAL = "Review theo quy trình được ban hành"

# --- Model GreenNode MaaS (OpenAI-compatible) để nhận diện hợp đồng trọng yếu ---
LLM_BASE_URL = os.environ.get("GREENNODE_BASE_URL", "https://maas-llm-aiplatform-hcm.api.vngcloud.vn/v1")
LLM_API_KEY  = os.environ.get("GREENNODE_API_KEY", "")
LLM_MODEL    = os.environ.get("GREENNODE_MODEL", "")  # để trống = tự dò model Qwen đang bật
_RESOLVED_MODEL = None

# DANH SÁCH HỢP ĐỒNG TRỌNG YẾU của công ty (tên + mô tả) — tiêu chí cho model nhận diện
MATERIAL_LIST = [
    ("Game licensing/sub-licensing", "HĐ với nhà phát triển (Tencent, Kingsoft, NetEase) để phát hành game theo khu vực."),
    ("IP Collaboration / Sub-licensing", "Dùng IP bên thứ ba trong game (nhân vật anime, Marvel, nhạc)."),
    ("Payment Gateway / Aggregator", "HĐ với ngân hàng, nhà mạng, ví điện tử (MoMo, ZaloPay) xử lý thanh toán trong game."),
    ("Platform distribution", "Thỏa thuận với Apple App Store, Google Play, Steam, PlayStation/Xbox."),
    ("Brand ambassador", "HĐ với người nổi tiếng/streamer lớn làm đại sứ thương hiệu."),
    ("Esport tournament / sponsorship", "Tổ chức giải đấu lớn hoặc tài trợ (Red Bull, Oppo)."),
    ("M&A và đầu tư", "Mua lại studio nhỏ hoặc nhận đầu tư."),
    ("Joint venture (liên doanh)", "Liên doanh với đối tác địa phương ở nước khác."),
    ("Banking & Credit Facility", "Khoản vay / hạn mức tín dụng tài trợ vận hành hoặc phí license."),
    ("Related party transactions", "HĐ giữa công ty với công ty mẹ hoặc công ty con (yêu cầu audit/compliance)."),
    ("IP Assignment", "Studio/freelancer/consultant tạo tài sản (art, code, nhạc) cho công ty."),
    ("Strategic NDA", "NDA liên quan game chưa công bố, M&A, hoặc chia sẻ source code."),
    ("Data Processing Agreement (DPA)", "Vendor chạm vào dữ liệu người dùng (ID, email, lịch sử thanh toán)."),
    ("Data Transfer Agreement", "Cho phép chuyển dữ liệu người chơi qua biên giới."),
    ("Exclusivity", "Cấm làm việc với đối thủ hoặc bắt buộc chỉ dùng một vendor."),
    ("Non-Compete", "Cam kết không vào một thị trường/loại game trong một thời gian."),
    ("Most Favored Nation (MFN)", "Cam kết không cho đối tác khác giá/điều khoản tốt hơn."),
    ("MoU / Letter of Intent (LOI)", "Dù ghi 'non-binding' vẫn thường có điều khoản ràng buộc (độc quyền, bảo mật, luật áp dụng)."),
    ("Settlement Agreement", "HĐ giải quyết tranh chấp/kiện tụng."),
    ("Platform Developer Terms", "Điều khoản chuẩn của Apple/Google/Steam để phát hành & kiếm tiền."),
    ("Unlimited Liability", "HĐ mà vendor không chịu giới hạn trách nhiệm của công ty."),
    ("Waiving Sovereign Immunity / đổi Governing Law", "Đưa công ty vào luật vùng tài phán rủi ro cao hoặc từ bỏ quyền kiện tại VN/Singapore."),
    ("Thời hạn ≥ 3 năm hoặc có auto-renew", "Bất kỳ HĐ nào ghi rõ thời hạn từ 3 năm trở lên hoặc có điều khoản tự động gia hạn."),
]

# Từ khóa fallback (chỉ dùng khi KHÔNG gọi được model) — xét ở tiêu đề + tên file
KEYWORDS_MATERIAL = [
    "non-disclosure","nondisclosure","nda","confidential","bảo mật","license","licens",
    "cấp phép","bản quyền","sub-licens","sublicens","intellectual property","sở hữu trí tuệ",
    "ip assignment","chuyển nhượng","payment gateway","cổng thanh toán","ví điện tử","momo","zalopay",
    "distribution","app store","google play","steam","playstation","xbox","ambassador","đại sứ",
    "esport","tournament","giải đấu","sponsorship","tài trợ","merger","acquisition","m&a","sáp nhập",
    "investment","đầu tư","joint venture","liên doanh","credit facility","khoản vay","tín dụng",
    "related party","data processing","dpa","data transfer","chuyển dữ liệu","exclusiv","độc quyền",
    "non-compete","most favored nation","mfn","memorandum of understanding","mou","letter of intent",
    "loi","settlement","unlimited liability","sovereign immunity","governing law","auto-renew","gia hạn",
]

_VN = "àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđ"

# ---------- Đọc văn bản nhiều định dạng ----------
def read_docx(path):
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for tb in d.tables:
        for row in tb.rows:
            parts.append(" ".join(c.text for c in row.cells))
    return "\n".join(parts)

def read_doc(path):
    """Chuyển .doc -> .docx bằng LibreOffice (nếu có) rồi đọc."""
    try:
        with tempfile.TemporaryDirectory() as tmp:
            subprocess.run(["soffice", "--headless", "--convert-to", "docx",
                            "--outdir", tmp, path], check=True,
                           stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            out = os.path.join(tmp, os.path.splitext(os.path.basename(path))[0] + ".docx")
            return read_docx(out) if os.path.exists(out) else ""
    except Exception:
        return ""   # môi trường không có LibreOffice -> bỏ qua file .doc

def read_pdf(path):
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join(pg.extract_text() or "" for pg in pdf.pages)
    except Exception:
        try:
            import PyPDF2
            r = PyPDF2.PdfReader(path)
            return "\n".join(pg.extract_text() or "" for pg in r.pages)
        except Exception:
            return ""

def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx": return read_docx(path)
    if ext == ".doc":  return read_doc(path)
    if ext == ".pdf":  return read_pdf(path)
    if ext in (".txt", ".md"):
        try:
            return open(path, encoding="utf-8", errors="ignore").read()
        except Exception:
            return ""
    return ""

SUPPORTED_EXT = (".docx", ".doc", ".pdf", ".txt", ".md")

# ---------- Chuẩn hóa & so sánh ----------
def normalize(text):
    """Bỏ qua khác biệt về số/ngày tháng, viết hoa, dấu câu, khoảng trắng."""
    t = text.lower()
    t = re.sub(r"\d+", " ", t)
    t = re.sub(rf"[^\w\s{_VN}]", " ", t, flags=re.UNICODE)
    return t.split()

def similarity(tokens_a, tokens_b):
    if not tokens_a or not tokens_b:
        return 0.0
    return difflib.SequenceMatcher(None, tokens_a, tokens_b).ratio() * 100.0

def first_line(text):
    for line in text.splitlines():
        if line.strip():
            return line.strip()
    return ""

def is_material(title, name):
    blob = (title + " " + name).lower()
    return any(k in blob for k in KEYWORDS_MATERIAL)

def _resolve_model():
    """Xác định mã model để gọi: ưu tiên biến môi trường, nếu trống thì tự dò model
    có chứa 'qwen' đang khả dụng qua /v1/models. Có cache để chỉ dò 1 lần."""
    global _RESOLVED_MODEL
    if LLM_MODEL:
        return LLM_MODEL
    if _RESOLVED_MODEL:
        return _RESOLVED_MODEL
    try:
        import requests
        r = requests.get(f"{LLM_BASE_URL}/models",
                         headers={"Authorization": f"Bearer {LLM_API_KEY}"}, timeout=15)
        r.raise_for_status()
        ids = [m.get("id", "") for m in r.json().get("data", [])]
        pick = next((i for i in ids if "qwen" in i.lower()), None) or (ids[0] if ids else None)
        _RESOLVED_MODEL = pick or "qwen/qwen3-5-27b"
    except Exception:
        _RESOLVED_MODEL = "qwen/qwen3-5-27b"
    return _RESOLVED_MODEL

def detect_material_llm(title, text):
    """Dùng model Qwen (GreenNode MaaS) xác định hợp đồng trọng yếu.
    Trả về dict {material, loai, ly_do} nếu gọi được; None nếu không (sẽ fallback dò từ khóa)."""
    if not LLM_API_KEY:
        return None
    try:
        import requests
        listing = "\n".join(f"- {n}: {d}" for n, d in MATERIAL_LIST)
        prompt = (
            "Bạn là trợ lý pháp chế của một công ty game. Dưới đây là DANH SÁCH các loại "
            "HỢP ĐỒNG TRỌNG YẾU của công ty (kèm mô tả):\n"
            f"{listing}\n\n"
            "Đọc đoạn đầu hợp đồng bên dưới và xác định nó có thuộc MỘT trong các loại trọng yếu "
            "ở trên không (dựa trên nội dung/bản chất, không chỉ tên gọi). Nếu có, ghi rõ thuộc loại nào.\n"
            'Chỉ trả lời bằng JSON: {"material": true/false, '
            '"loai": "<tên loại theo danh sách, hoặc loại hợp đồng nếu không trọng yếu>", '
            '"ly_do": "<1 câu ngắn giải thích>"}.\n\n'
            f"Tiêu đề: {title}\nNội dung (rút gọn):\n{text[:3500]}"
        )
        r = requests.post(
            f"{LLM_BASE_URL}/chat/completions",
            headers={"Authorization": f"Bearer {LLM_API_KEY}", "Content-Type": "application/json"},
            json={"model": _resolve_model(),
                  "messages": [{"role": "user", "content": prompt}],
                  "temperature": 0},
            timeout=30,
        )
        r.raise_for_status()
        content = r.json()["choices"][0]["message"]["content"]
        m = re.search(r"\{.*\}", content, re.S)
        data = json.loads(m.group(0)) if m else {}
        return {"material": bool(data.get("material")),
                "loai": str(data.get("loai", "")),
                "ly_do": str(data.get("ly_do", ""))}
    except Exception:
        return None

# ---------- API cấp cao ----------
def build_templates(template_paths):
    """Nhận list đường dẫn file template -> list dict {name, tokens}."""
    templates = []
    for p in template_paths:
        txt = extract_text(p)
        if txt.strip():
            templates.append({"name": os.path.basename(p), "tokens": normalize(txt)})
    return templates

def classify_text(name, text, templates, threshold=THRESHOLD):
    """Phân loại 1 hợp đồng. Template TÙY CHỌN:
       - Có template: phân theo % giống (logic cũ) + đánh dấu trọng yếu.
       - Không template: chỉ phân trọng yếu / không trọng yếu.
    Mọi hợp đồng đều được gắn nhãn trọng yếu/không và 1 ghi chú quy trình (note)."""
    tokens = normalize(text)
    title = first_line(text)
    best_name, best_sim = "-", 0.0
    for tpl in templates:
        s = similarity(tokens, tpl["tokens"])
        if s > best_sim:
            best_sim, best_name = s, tpl["name"]

    # Luôn xác định trọng yếu theo DANH SÁCH (ưu tiên Qwen, fallback dò từ khóa)
    llm = detect_material_llm(title, text)
    if llm is not None:
        material, loai, ly_do = llm["material"], llm["loai"], llm["ly_do"]
    else:
        material, loai, ly_do = is_material(title, name), "", ""

    note = NOTE_MATERIAL if material else NOTE_NONMATERIAL

    if templates:                       # có template -> giữ nguyên cách phân theo template
        if best_sim >= threshold:
            cat = CAT_SIM
        elif material:
            cat = CAT_KEY
        else:
            cat = CAT_NONE
    else:                               # không template -> chỉ trọng yếu / không trọng yếu
        cat = CAT_KEY if material else CAT_NONMATERIAL

    return {"name": name, "title": title[:90], "best_tpl": best_name,
            "sim": round(best_sim, 1), "material": material,
            "loai": loai, "ly_do": ly_do, "note": note, "cat": cat}

def classify_file(path, templates, threshold=THRESHOLD, display_name=None):
    text = extract_text(path)
    return classify_text(display_name or os.path.basename(path), text, templates, threshold)
